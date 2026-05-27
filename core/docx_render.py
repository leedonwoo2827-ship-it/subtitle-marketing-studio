"""DOCX rendering for press release (Word file for journalist email distribution)."""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path


@dataclass
class DOCXRenderResult:
    path: Path | None
    error: str = ""

    @property
    def ok(self) -> bool:
        return self.path is not None and not self.error


def _set_korean_font(run, *, font: str = "맑은 고딕", size_pt: float = 10.0) -> None:
    """Set East-Asian font + size on a python-docx run."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size_pt)


def _strip_md_marks(text: str) -> str:
    """Strip **bold** / *italic* / `code` markers — keep content only."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def press_release(md_text: str, out_dir: Path, *, filename: str = "output.docx") -> DOCXRenderResult:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return DOCXRenderResult(path=None, error="python-docx not installed. Run: pip install python-docx")

    try:
        doc = Document()

        # Set default paragraph spacing
        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "맑은 고딕"
        normal.font.size = Pt(10)

        lines = md_text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line:
                i += 1
                continue
            if line.startswith("# "):
                # Headline
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(_strip_md_marks(line[2:].strip("[]")))
                run.bold = True
                _set_korean_font(run, size_pt=18)
                p.paragraph_format.space_after = Pt(8)
            elif line.startswith("## "):
                # Section header
                text = _strip_md_marks(line[3:])
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                _set_korean_font(run, size_pt=14)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
            elif line.startswith("### "):
                # Sub-section
                text = _strip_md_marks(line[4:])
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                _set_korean_font(run, size_pt=12)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
            elif re.match(r"^\s*[-*]\s+", line):
                # Bullet
                text = _strip_md_marks(re.sub(r"^\s*[-*]\s+", "", line))
                p = doc.add_paragraph(style="List Bullet")
                run = p.runs[0] if p.runs else p.add_run()
                if not p.runs:
                    p.add_run(text)
                else:
                    p.runs[0].text = text
                for run in p.runs:
                    _set_korean_font(run, size_pt=10)
            elif line.startswith("> "):
                # Blockquote
                p = doc.add_paragraph()
                run = p.add_run(_strip_md_marks(line[2:]))
                run.italic = True
                _set_korean_font(run, size_pt=10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                # Regular paragraph
                p = doc.add_paragraph()
                run = p.add_run(_strip_md_marks(line))
                _set_korean_font(run, size_pt=10)
                p.paragraph_format.line_spacing = 1.5
            i += 1

        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / filename
        doc.save(str(out_path))
        return DOCXRenderResult(path=out_path)
    except Exception as e:
        return DOCXRenderResult(path=None, error=f"{type(e).__name__}: {e}")


RENDERERS: dict[str, callable] = {
    "press_release": press_release,
}


def render(renderer_key: str, md_text: str, out_dir: Path) -> DOCXRenderResult:
    fn = RENDERERS.get(renderer_key)
    if fn is None:
        return DOCXRenderResult(path=None, error=f"unknown renderer: {renderer_key}")
    return fn(md_text, out_dir)
